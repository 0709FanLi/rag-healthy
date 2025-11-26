import os
import io
import asyncio
from typing import List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pypdf
import docx
import openpyxl

from src.models.tables.knowledge import KnowledgeFile, FileStatus
from src.services.oss_service import OSSService
from src.services.volc_service import VolcService
from src.services.vector_service import VectorService

class KnowledgeService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.oss = OSSService()
        self.volc = VolcService()
        self.vector = VectorService()

    async def process_file_background(self, file_id: int):
        """后台处理文件：下载 -> 解析 -> 向量化 -> 存储"""
        try:
            # 1. 获取文件记录
            result = await self.db.execute(select(KnowledgeFile).where(KnowledgeFile.id == file_id))
            file_record = result.scalar_one_or_none()
            if not file_record:
                return

            # 更新状态为处理中
            file_record.status = FileStatus.processing
            await self.db.commit()

            # 2. 下载文件
            content = self.oss.download_file(file_record.oss_url)
            
            # 3. 提取文本
            text_content = self._extract_text(content, file_record.filename)
            
            if not text_content.strip():
                raise Exception("Empty text content extracted")

            # 4. 文本切片
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50,
                separators=["\n\n", "\n", "。", "！", "？", " ", ""]
            )
            chunks = splitter.split_text(text_content)

            # 5. 向量化 (分批处理)
            batch_size = 10
            vectors = []
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i+batch_size]
                batch_vectors = await self.volc.get_embeddings(batch)
                vectors.extend(batch_vectors)

            # 6. 存储到 Weaviate
            weaviate_chunks = []
            for i, chunk in enumerate(chunks):
                weaviate_chunks.append({
                    "content": chunk,
                    "doc_id": file_record.id,
                    "kb_type": file_record.kb_type.value,
                    "tags": file_record.tags,
                    "source": file_record.filename,
                    "vector": vectors[i]
                })
            
            # 确保 Schema 存在
            self.vector.init_schema()
            self.vector.add_chunks(weaviate_chunks)

            # 7. 更新状态为完成
            file_record.status = FileStatus.completed
            await self.db.commit()

        except Exception as e:
            print(f"Error processing file {file_id}: {e}")
            # 重新获取 session 避免 transaction 错误
            # 注意：在实际生产中，可能需要更好的错误恢复机制
            file_record.status = FileStatus.failed
            file_record.error_msg = str(e)
            await self.db.commit()

    def _extract_text(self, content: bytes, filename: str) -> str:
        """根据文件扩展名提取文本."""
        ext = os.path.splitext(filename)[1].lower()
        file_stream = io.BytesIO(content)

        if ext == '.pdf':
            reader = pypdf.PdfReader(file_stream)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif ext in ['.docx', '.doc']:
            doc = docx.Document(file_stream)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext == '.txt':
            return content.decode('utf-8')
        elif ext in ['.xlsx', '.xls']:
            wb = openpyxl.load_workbook(file_stream)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.values:
                    text += " ".join([str(cell) for cell in row if cell]) + "\n"
            return text
        else:
            raise Exception(f"Unsupported file type: {ext}")

